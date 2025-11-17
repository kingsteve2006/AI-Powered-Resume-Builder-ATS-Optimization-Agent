#!/usr/bin/env python3
"""
AI-Powered Resume Builder & ATS Optimization Agent
Streamlit Web Application with OpenAI and Gemini API Integration
"""

import streamlit as st
import openai
import google.generativeai as genai
import json
import os
from typing import Dict, List, Any, Optional
from datetime import datetime
import tempfile
import uuid
from pathlib import Path
from io import BytesIO
import random

# Document generation libraries
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.pdfgen import canvas
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# File processing libraries
try:
    import PyPDF2
    import docx2txt
    FILE_PROCESSING_AVAILABLE = True
except ImportError:
    FILE_PROCESSING_AVAILABLE = False

# Configure APIs
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "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")
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "AIzaSyBtKLSxlo5sWi2vuasBsleLJHYpt5KumKc")

# Initialize OpenAI client
client = openai.OpenAI(api_key=OPENAI_API_KEY)

# Initialize Gemini
genai.configure(api_key=GEMINI_API_KEY)
gemini_model = genai.GenerativeModel('gemini-1.5-flash-latest')

class ResumeData:
    """Data class to hold resume information"""
    def __init__(self):
        self.personal = {}
        self.summary = ""
        self.education = []
        self.skills = ""
        self.experience = []
        self.projects = []

class ATSScorer:
    """ATS Scoring System"""
    
    @staticmethod
    def calculate_score(resume_data: ResumeData) -> Dict[str, Any]:
        """Calculate ATS score and provide detailed breakdown"""
        contact_score = 0
        experience_score = 0
        education_score = 0
        skills_score = 0
        
        # Contact information (0-20 points)
        if resume_data.personal.get("name"): contact_score += 4
        if resume_data.personal.get("email"): contact_score += 4
        if resume_data.personal.get("phone"): contact_score += 4
        if resume_data.personal.get("location"): contact_score += 4
        if resume_data.personal.get("linkedin"): contact_score += 2
        if resume_data.personal.get("portfolio"): contact_score += 2
        
        # Work experience (0-30 points)
        if resume_data.experience:
            experience_score += 10
            for exp in resume_data.experience:
                if exp.get("title") and exp.get("company"): experience_score += 5
                if exp.get("description") and len(exp.get("description", "")) > 50: experience_score += 5
        
        # Education (0-20 points)
        if resume_data.education:
            education_score += 10
            for edu in resume_data.education:
                if edu.get("degree") and edu.get("institution"): education_score += 5
        
        # Skills (0-30 points)
        if resume_data.skills:
            skill_count = len([s.strip() for s in resume_data.skills.split(",") if s.strip()])
            skills_score = min(skill_count * 3, 30)
        
        total_score = contact_score + experience_score + education_score + skills_score
        
        return {
            "total_score": min(total_score, 100),
            "contact": contact_score,
            "experience": experience_score,
            "education": education_score,
            "skills": skills_score,
            "status": "Excellent" if total_score >= 80 else "Good" if total_score >= 60 else "Needs Improvement"
        }

class AIEnhancer:
    """AI Enhancement using OpenAI and Gemini APIs"""
    
    @staticmethod
    def call_openai(prompt: str, max_tokens: int = 500) -> Optional[str]:
        """Call OpenAI API with error handling"""
        try:
            response = client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=max_tokens,
                temperature=0.7
            )
            return response.choices[0].message.content.strip()
        except Exception as e:
            st.error(f"OpenAI API Error: {str(e)}")
            return None
    
    @staticmethod
    def call_gemini(prompt: str) -> Optional[str]:
        """Call Gemini API with error handling"""
        try:
            response = gemini_model.generate_content(prompt)
            return response.text.strip()
        except Exception as e:
            st.error(f"Gemini API Error: {str(e)}")
            return None
    
    @staticmethod
    def enhance_summary(summary: str) -> str:
        """Enhance professional summary using AI"""
        prompt = f"""Enhance this professional summary for a resume to make it more compelling and ATS-friendly:

Original: "{summary}"

Provide an enhanced version that:
1. Uses action verbs and quantifiable achievements
2. Includes relevant industry keywords
3. Maintains professional tone
4. Is optimized for ATS systems

Enhanced summary:"""
        
        enhanced = AIEnhancer.call_openai(prompt, 200)
        if not enhanced:
            enhanced = AIEnhancer.call_gemini(prompt)
        
        if not enhanced:
            return AIEnhancer.fallback_summary_enhancement(summary)
        
        return enhanced
    
    @staticmethod
    def enhance_description(description: str) -> str:
        """Enhance job description using AI"""
        if not description:
            return ""
            
        prompt = f"""Enhance this job description for a resume:

Original: "{description}"

Provide an enhanced version with:
1. Strong action verbs
2. Specific metrics and numbers
3. Industry-relevant keywords
4. Concise but impactful language

Enhanced description:"""
        
        enhanced = AIEnhancer.call_openai(prompt, 300)
        if not enhanced:
            enhanced = AIEnhancer.call_gemini(prompt)
        
        if not enhanced:
            return AIEnhancer.fallback_description_enhancement(description)
        
        return enhanced
    
    @staticmethod
    def enhance_skills(skills: str) -> str:
        """Enhance skills section using AI"""
        prompt = f"""Enhance this skills section for a resume:

Original: "{skills}"

Provide an enhanced version that:
1. Organizes skills by category
2. Includes current industry technologies
3. Uses ATS-friendly keywords

Enhanced skills:"""
        
        enhanced = AIEnhancer.call_openai(prompt, 400)
        if not enhanced:
            enhanced = AIEnhancer.call_gemini(prompt)
        
        if not enhanced:
            return AIEnhancer.fallback_skills_enhancement(skills)
        
        return enhanced
    
    @staticmethod
    def fallback_summary_enhancement(summary: str) -> str:
        """Fallback summary enhancement"""
        enhancements = [
            "Results-driven professional with extensive experience in cross-functional team leadership and strategic project execution. Proven track record of delivering innovative solutions that drive business growth and operational efficiency.",
            "Accomplished specialist with deep expertise in modern technologies and best practices. Demonstrated ability to optimize processes, mentor teams, and deliver high-impact projects that exceed stakeholder expectations.",
            "Dynamic professional combining technical expertise with strong leadership skills. Passionate about leveraging emerging technologies to solve complex challenges and create sustainable competitive advantages."
        ]
        return random.choice(enhancements)
    
    @staticmethod
    def fallback_description_enhancement(description: str) -> str:
        """Fallback description enhancement"""
        enhancements = [
            "• Led cross-functional team of 8 developers in delivering mission-critical web application, resulting in 40% improvement in user engagement and $200K annual cost savings",
            "• Architected and implemented scalable microservices infrastructure using modern frameworks, handling 100K+ daily active users with 99.9% uptime",
            "• Spearheaded digital transformation initiative, modernizing legacy systems and reducing technical debt by 60% while improving system performance by 3x"
        ]
        return description + "\n" + random.choice(enhancements) if description else random.choice(enhancements)
    
    @staticmethod
    def fallback_skills_enhancement(skills: str) -> str:
        """Fallback skills enhancement"""
        skill_enhancements = {
            "javascript": "JavaScript (ES6+, TypeScript, Node.js, React, Vue.js)",
            "python": "Python (Django, Flask, FastAPI, Data Science, Machine Learning)",
            "java": "Java (Spring Boot, Microservices, Enterprise Applications)",
            "react": "React (Hooks, Context API, Redux, Next.js, Testing Library)",
            "html": "HTML5, CSS3, Responsive Design, Bootstrap, Tailwind CSS",
            "aws": "AWS (EC2, S3, Lambda, RDS, CloudFormation)",
            "docker": "Docker, Kubernetes, Container Orchestration"
        }
        
        enhanced = skills
        for key, value in skill_enhancements.items():
            if key.lower() in skills.lower():
                enhanced = enhanced.replace(key, value)
        
        return enhanced

class DocumentGenerator:
    """Generate PDF and Word documents"""
    
    @staticmethod
    def generate_pdf(resume_data: ResumeData) -> BytesIO:
        """Generate PDF resume"""
        if not PDF_AVAILABLE:
            raise Exception("PDF generation not available. Install reportlab.")
        
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Title
        title = Paragraph(f"<b>{resume_data.personal.get('name', 'Your Name')}</b>", styles['Title'])
        story.append(title)
        story.append(Spacer(1, 12))
        
        # Contact info
        contact = f"{resume_data.personal.get('email', '')} | {resume_data.personal.get('phone', '')}"
        if resume_data.personal.get('location'):
            contact += f" | {resume_data.personal.get('location')}"
        story.append(Paragraph(contact, styles['Normal']))
        story.append(Spacer(1, 12))
        
        # Summary
        if resume_data.summary:
            story.append(Paragraph("<b>Professional Summary</b>", styles['Heading2']))
            story.append(Paragraph(resume_data.summary, styles['Normal']))
            story.append(Spacer(1, 12))
        
        # Experience
        if resume_data.experience:
            story.append(Paragraph("<b>Work Experience</b>", styles['Heading2']))
            for exp in resume_data.experience:
                story.append(Paragraph(f"<b>{exp.get('title', '')}</b> - {exp.get('company', '')}", styles['Heading3']))
                story.append(Paragraph(f"{exp.get('startDate', '')} - {exp.get('endDate', '')}", styles['Normal']))
                if exp.get('description'):
                    story.append(Paragraph(exp['description'], styles['Normal']))
                story.append(Spacer(1, 6))
        
        # Skills
        if resume_data.skills:
            story.append(Paragraph("<b>Skills</b>", styles['Heading2']))
            story.append(Paragraph(resume_data.skills, styles['Normal']))
        
        doc.build(story)
        buffer.seek(0)
        return buffer
    
    @staticmethod
    def generate_word(resume_data: ResumeData) -> BytesIO:
        """Generate Word resume"""
        if not DOCX_AVAILABLE:
            raise Exception("Word generation not available. Install python-docx.")
        
        doc = Document()
        
        # Title
        title = doc.add_heading(resume_data.personal.get('name', 'Your Name'), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Contact
        contact = f"{resume_data.personal.get('email', '')} | {resume_data.personal.get('phone', '')}"
        if resume_data.personal.get('location'):
            contact += f" | {resume_data.personal.get('location')}"
        p = doc.add_paragraph(contact)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Summary
        if resume_data.summary:
            doc.add_heading('Professional Summary', level=1)
            doc.add_paragraph(resume_data.summary)
        
        # Experience
        if resume_data.experience:
            doc.add_heading('Work Experience', level=1)
            for exp in resume_data.experience:
                doc.add_heading(f"{exp.get('title', '')} - {exp.get('company', '')}", level=2)
                doc.add_paragraph(f"{exp.get('startDate', '')} - {exp.get('endDate', '')}")
                if exp.get('description'):
                    doc.add_paragraph(exp['description'])
        
        # Skills
        if resume_data.skills:
            doc.add_heading('Skills', level=1)
            doc.add_paragraph(resume_data.skills)
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

def main():
    """Main Streamlit application"""
    st.set_page_config(
        page_title="AI Resume Builder",
        page_icon="📄",
        layout="wide"
    )
    
    # Custom CSS
    st.markdown("""
    <style>
    .main-header {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 2rem;
        border-radius: 10px;
        color: white;
        text-align: center;
        margin-bottom: 2rem;
    }
    .score-card {
        background: #f8f9fa;
        padding: 1rem;
        border-radius: 10px;
        border-left: 4px solid #007bff;
        margin: 1rem 0;
    }
    </style>
    """, unsafe_allow_html=True)
    
    # Header
    st.markdown("""
    <div class="main-header">
        <h1>🚀 AI-Powered Resume Builder</h1>
        <p>ATS Optimization & Professional Resume Generation</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Initialize session state
    if 'resume_data' not in st.session_state:
        st.session_state.resume_data = ResumeData()
    if 'enhanced_data' not in st.session_state:
        st.session_state.enhanced_data = None
    if 'ats_score' not in st.session_state:
        st.session_state.ats_score = None
    if 'current_step' not in st.session_state:
        st.session_state.current_step = 1
    
    # Sidebar navigation
    st.sidebar.title("Navigation")
    steps = [
        "📝 Input Resume Data",
        "📊 ATS Score Analysis", 
        "🤖 AI Enhancement",
        "📄 Generate Documents"
    ]
    
    selected_step = st.sidebar.radio("Choose Step:", steps, index=st.session_state.current_step-1)
    st.session_state.current_step = steps.index(selected_step) + 1
    
    # Step 1: Input Resume Data
    if st.session_state.current_step == 1:
        st.header("Step 1: Input Resume Data")
        
        # Personal Information
        st.subheader("Personal Information")
        col1, col2 = st.columns(2)
        
        with col1:
            st.session_state.resume_data.personal['name'] = st.text_input("Full Name", value=st.session_state.resume_data.personal.get('name', ''))
            st.session_state.resume_data.personal['email'] = st.text_input("Email", value=st.session_state.resume_data.personal.get('email', ''))
            st.session_state.resume_data.personal['phone'] = st.text_input("Phone", value=st.session_state.resume_data.personal.get('phone', ''))
        
        with col2:
            st.session_state.resume_data.personal['location'] = st.text_input("Location", value=st.session_state.resume_data.personal.get('location', ''))
            st.session_state.resume_data.personal['linkedin'] = st.text_input("LinkedIn", value=st.session_state.resume_data.personal.get('linkedin', ''))
            st.session_state.resume_data.personal['portfolio'] = st.text_input("Portfolio", value=st.session_state.resume_data.personal.get('portfolio', ''))
        
        st.subheader("Professional Summary")
        st.session_state.resume_data.summary = st.text_area("Summary", value=st.session_state.resume_data.summary, height=100)
        
        st.subheader("Skills")
        st.session_state.resume_data.skills = st.text_area("Skills (comma-separated)", value=st.session_state.resume_data.skills, height=100)
        
        st.subheader("Work Experience")
        if 'experience_count' not in st.session_state:
            st.session_state.experience_count = max(1, len(st.session_state.resume_data.experience))
        
        for i in range(st.session_state.experience_count):
            with st.expander(f"Experience {i+1}", expanded=True):
                if i >= len(st.session_state.resume_data.experience):
                    st.session_state.resume_data.experience.append({})
                
                col1, col2 = st.columns(2)
                with col1:
                    st.session_state.resume_data.experience[i]['title'] = st.text_input(f"Job Title {i+1}", value=st.session_state.resume_data.experience[i].get('title', ''), key=f"title_{i}")
                    st.session_state.resume_data.experience[i]['company'] = st.text_input(f"Company {i+1}", value=st.session_state.resume_data.experience[i].get('company', ''), key=f"company_{i}")
                
                with col2:
                    st.session_state.resume_data.experience[i]['startDate'] = st.text_input(f"Start Date {i+1}", value=st.session_state.resume_data.experience[i].get('startDate', ''), key=f"start_{i}")
                    st.session_state.resume_data.experience[i]['endDate'] = st.text_input(f"End Date {i+1}", value=st.session_state.resume_data.experience[i].get('endDate', ''), key=f"end_{i}")
                
                st.session_state.resume_data.experience[i]['description'] = st.text_area(f"Description {i+1}", value=st.session_state.resume_data.experience[i].get('description', ''), key=f"desc_{i}")
        
        col1, col2 = st.columns(2)
        with col1:
            if st.button("Add Experience"):
                st.session_state.experience_count += 1
                st.rerun()
        
        with col2:
            if st.button("Remove Last Experience") and st.session_state.experience_count > 1:
                st.session_state.experience_count -= 1
                if len(st.session_state.resume_data.experience) > st.session_state.experience_count:
                    st.session_state.resume_data.experience.pop()
                st.rerun()
        
        if st.button("Calculate ATS Score", type="primary"):
            st.session_state.current_step = 2
            st.rerun()
    
    # Step 2: ATS Score Analysis
    elif st.session_state.current_step == 2:
        st.header("Step 2: ATS Score Analysis")
        
        # Calculate ATS score
        st.session_state.ats_score = ATSScorer.calculate_score(st.session_state.resume_data)
        score_data = st.session_state.ats_score
        
        col1, col2 = st.columns([1, 2])
        
        with col1:
            # Score display
            st.markdown(f"""
            <div class="score-card">
                <h2 style="text-align: center; color: {'#28a745' if score_data['total_score'] >= 80 else '#ffc107' if score_data['total_score'] >= 60 else '#dc3545'}">
                    {score_data['total_score']}/100
                </h2>
                <h4 style="text-align: center;">{score_data['status']}</h4>
            </div>
            """, unsafe_allow_html=True)
        
        with col2:
            # Score breakdown
            st.subheader("Score Breakdown")
            
            categories = [
                ("Contact Information", score_data['contact'], 20),
                ("Work Experience", score_data['experience'], 30),
                ("Education", score_data['education'], 20),
                ("Skills & Keywords", score_data['skills'], 30)
            ]
            
            for category, score, max_score in categories:
                percentage = (score / max_score) * 100 if max_score > 0 else 0
                st.write(f"**{category}**: {score}/{max_score}")
                st.progress(percentage / 100)
        
        if st.button("Enhance with AI", type="primary"):
            st.session_state.current_step = 3
            st.rerun()
    
    # Step 3: AI Enhancement
    elif st.session_state.current_step == 3:
        st.header("Step 3: AI Enhancement")
        
        if st.session_state.enhanced_data is None:
            with st.spinner("AI is enhancing your resume..."):
                # Create enhanced copy
                enhanced = ResumeData()
                enhanced.personal = st.session_state.resume_data.personal.copy()
                enhanced.education = st.session_state.resume_data.education.copy()
                enhanced.projects = st.session_state.resume_data.projects.copy()
                
                # Enhance summary
                if st.session_state.resume_data.summary:
                    enhanced.summary = AIEnhancer.enhance_summary(st.session_state.resume_data.summary)
                else:
                    enhanced.summary = st.session_state.resume_data.summary
                
                # Enhance experience
                enhanced.experience = []
                for exp in st.session_state.resume_data.experience:
                    enhanced_exp = exp.copy()
                    if exp.get('description'):
                        enhanced_exp['description'] = AIEnhancer.enhance_description(exp['description'])
                    enhanced.experience.append(enhanced_exp)
                
                # Enhance skills
                if st.session_state.resume_data.skills:
                    enhanced.skills = AIEnhancer.enhance_skills(st.session_state.resume_data.skills)
                else:
                    enhanced.skills = st.session_state.resume_data.skills
                
                st.session_state.enhanced_data = enhanced
        
        # Show comparison
        st.subheader("Before vs After Enhancement")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("**Original**")
            st.write(f"**Summary:** {st.session_state.resume_data.summary[:200]}..." if len(st.session_state.resume_data.summary) > 200 else f"**Summary:** {st.session_state.resume_data.summary}")
            if st.session_state.resume_data.experience:
                desc = st.session_state.resume_data.experience[0].get('description', '')
                st.write(f"**First Job Description:** {desc[:200]}..." if len(desc) > 200 else f"**First Job Description:** {desc}")
        
        with col2:
            st.markdown("**Enhanced**")
            st.write(f"**Summary:** {st.session_state.enhanced_data.summary[:200]}..." if len(st.session_state.enhanced_data.summary) > 200 else f"**Summary:** {st.session_state.enhanced_data.summary}")
            if st.session_state.enhanced_data.experience:
                desc = st.session_state.enhanced_data.experience[0].get('description', '')
                st.write(f"**First Job Description:** {desc[:200]}..." if len(desc) > 200 else f"**First Job Description:** {desc}")
        
        # Score improvement
        original_score = st.session_state.ats_score['total_score']
        enhanced_score_data = ATSScorer.calculate_score(st.session_state.enhanced_data)
        enhanced_score = enhanced_score_data['total_score']
        improvement = enhanced_score - original_score
        
        st.success(f"Score improved from {original_score} to {enhanced_score} (+{improvement} points)")
        
        if st.button("Generate Documents", type="primary"):
            st.session_state.current_step = 4
            st.rerun()
    
    # Step 4: Generate Documents
    elif st.session_state.current_step == 4:
        st.header("Step 4: Generate Documents")
        
        if st.session_state.enhanced_data is None:
            st.error("Please complete the AI enhancement step first.")
            return
        
        col1, col2 = st.columns(2)
        
        with col1:
            if st.button("Generate PDF", type="primary"):
                try:
                    with st.spinner("Generating PDF..."):
                        pdf_buffer = DocumentGenerator.generate_pdf(st.session_state.enhanced_data)
                        
                    st.download_button(
                        label="Download PDF",
                        data=pdf_buffer.getvalue(),
                        file_name=f"resume_{st.session_state.enhanced_data.personal.get('name', 'resume').replace(' ', '_')}.pdf",
                        mime="application/pdf"
                    )
                    st.success("PDF generated successfully!")
                except Exception as e:
                    st.error(f"Error generating PDF: {str(e)}")
        
        with col2:
            if st.button("Generate Word", type="primary"):
                try:
                    with st.spinner("Generating Word document..."):
                        word_buffer = DocumentGenerator.generate_word(st.session_state.enhanced_data)
                    
                    st.download_button(
                        label="Download Word",
                        data=word_buffer.getvalue(),
                        file_name=f"resume_{st.session_state.enhanced_data.personal.get('name', 'resume').replace(' ', '_')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    st.success("Word document generated successfully!")
                except Exception as e:
                    st.error(f"Error generating Word document: {str(e)}")
        
        # Preview
        st.subheader("Resume Preview")
        with st.expander("View Enhanced Resume", expanded=True):
            st.write(f"**Name:** {st.session_state.enhanced_data.personal.get('name', '')}")
            st.write(f"**Email:** {st.session_state.enhanced_data.personal.get('email', '')}")
            st.write(f"**Phone:** {st.session_state.enhanced_data.personal.get('phone', '')}")
            
            if st.session_state.enhanced_data.summary:
                st.write(f"**Summary:** {st.session_state.enhanced_data.summary}")
            
            if st.session_state.enhanced_data.experience:
                st.write("**Experience:**")
                for i, exp in enumerate(st.session_state.enhanced_data.experience):
                    st.write(f"{i+1}. **{exp.get('title', '')}** at {exp.get('company', '')}")
                    st.write(f"   {exp.get('startDate', '')} - {exp.get('endDate', '')}")
                    if exp.get('description'):
                        st.write(f"   {exp['description']}")
            
            if st.session_state.enhanced_data.skills:
                st.write(f"**Skills:** {st.session_state.enhanced_data.skills}")

if __name__ == "__main__":
    main()